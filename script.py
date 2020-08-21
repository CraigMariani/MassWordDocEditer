'''
Requirements for runnning script

 - Python3
 - git bash command line 
 - docx python module (type following command into command line to install: pip install python-docx)
 - subprocess python module (type following command into command line to install: pip install subprocess)

 - This script will only work with files of the .docx extension so be sure to save files like this 
 - Script will also need to be in the same directory as the files you are trying to change
 - The script iterates through all files in the directory that it is in and saves it in a new folder called "results_documents"
 - This folder is automitcally created if it does not exist

'''

import os 
import docx
import subprocess
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


global files_unchanged
global files_changed

global old_value
global old_value2

global new_value
global new_value2
class Edit_Document:

    # constructor, initializes variables
    def __init__(self, file_name):
        doc = docx.Document(file_name) # create document object
        self.doc = doc
        self.file_name = file_name

    # also sets formatting but for footers
    def set_footer_style(self):
        footer_style = self.doc.styles['Normal']
        font = footer_style.font
        font.name = 'Arial'
        font.size = Pt(8)

        return footer_style
    # for setting the formatting that you want when you make changes
    def set_style(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(14)

        return style

    # change the foot notes
    def change_footers(self):
        doc = self.doc
        # iterate through footer sections on each page
        for page_num, section in enumerate(doc.sections):
            # footer = section.footer()
            # print(footer.paragraphs[0].text)
            footer_style = ed.set_footer_style()

            try:
                # change footers on first page
                for footer_line in section.first_page_footer.paragraphs:
                    footer_line
                    footer_line.text = footer_line.text.replace(old_value, new_value)
                    footer_line.text = footer_line.text.replace(old_value2, new_value2)
                
                # change footers in odd pages
                for footer_line in section.footer.paragraphs:
                    
                    footer_line.text = footer_line.text.replace(old_Value, new_value)
                    footer_line.text = footer_line.text.replace(old_value2, new_value2)
                
                # change footers in even pages
                for footer_line in section.even_page_footer.paragraphs:
                    footer_line.text = footer_line.text.replace(old_value, new_value)
                    footer_line.text = footer_line.text.replace(old_value2, new_value2)
                
                print('footers changed')
            except Exception as e:
                print('Could not change footers for document{}'.format(file_name))


        try:
            
            path = 'results_documents/{}'.format(file_name)
            doc.save(path)
            print('file: {} saved successfully'.format(file_name))

        except FileNotFoundError:
            print('File not found, creating file')
            subprocess.call(['mkdir', 'results_documents'])
            path = 'results_documents/{}'.format(file_name)
            doc.save(path)
            print('file: {} saved successfully'.format(file_name))
        
        except  PermissionError:
            print('not able to get permsissions for file: {}'.format(file_name))
            files_unchanged.append(file_name)

    # change the text in main document
    def change_text(self):
        doc = self.doc
        file_name = self.file_name

        lines = list(doc.paragraphs)


        style = ed.set_style()

        for index, line in enumerate(lines):
            
            
            
            if old_value == line.text: # <<<< where we look for the old value we want to change
                line.style = doc.styles['Normal']
                line.alignment = WD_ALIGN_PARAGRAPH.CENTER
                line.bold = True
                line.text = new_value # the new value that we are changing it to
                
                print('Change made: {} at line: {}'.format(line.text, index))
            
            
            if old_value2 == line.text: # <<<< where we look for the old value we want to change
                line.style = doc.styles['Normal']
                line.alignment = WD_ALIGN_PARAGRAPH.CENTER
                line.bold = True
                line.text = new_value2 # the new value that we are changing it to

                print('Change made: {} at line: {}'.format(line.text, index))
        

        try:
            
            path = 'results_documents/{}'.format(file_name)
            doc.save(path)
            print('file: {} saved successfully'.format(file_name))

        except FileNotFoundError:
            print('Folder: results_documents not found, creating Folder')
            subprocess.call(['mkdir', 'results_documents'])
            path = 'results_documents/{}'.format(file_name)
            doc.save(path)
            print('file: {} saved successfully'.format(file_name))
        
        except  PermissionError:
            print('not able to get permsissions for file: {}'.format(file_name))
            files_unchanged.append(file_name)


    # display line count of document and  text inside document
    def show_text(self):
        doc = self.doc
        line_count = len(list(doc.paragraphs))

        lines = list(doc.paragraphs)
        print('LINE COUNT {}'.format(line_count))

        for line in lines:
            print(line.text)


    
if __name__ == '__main__':
    
    files = os.listdir()

    files_unchanged = []
    files_changed = []

    print('input two string values seperated by spaces in the following format')
    print('{old value} {new value} {old_value2} {new_value2}')
    old_value, new_value, old_value2, new_value2 = map(str, input().split())


    for file_name in files:

        if file_name == 'script.py':
            continue

        if '.docx' not in file_name:
            print('Could not edit file {}'.format(file_name))
            print('Please save the file as .docx extension')
            files_unchanged.append(file_name)
            continue
            
        ed = Edit_Document(file_name)
        ed.change_text()
        ed.change_footers()
        files_changed.append(file_name)
        
    print()
    print('\/*\/*\/*\/*\/*\/*\/*\/*\/*\/')
    print('\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/')
    print('\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/')
    print()
    print('files edited: {}'.format(files_changed))
    print()
    print('files unedited: {}'.format(files_unchanged))
    print()
    print('\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/')
    print('\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/*\/')
    print('\/*\/*\/*\/*\/*\/*\/*\/*\/*\/')